"""
Provides small I/O helpers: pick the latest csv or parquet dump, write new
parquet snapshots, register seaborn palettes, and save dataframes to Word.
"""
from collections import Counter
from datetime import datetime, UTC
from DB import engine
from DB.ops import load_scrobble_table_from_db_to_df
from docx import Document
import pandas as pd
from pathlib import Path
import re
if '__file__' in globals():
    HERE = Path(__file__).resolve().parent
    # Get the project root directory (parent of the helpers directory)
    PROJECT_ROOT = HERE.parent
else:
    PROJECT_ROOT = Path.cwd()
PQ_DIR = PROJECT_ROOT / "PQ"
PQ_DIR.mkdir(exist_ok=True)
OP_TOKENS = {           # space–operator–space → token
    r"\s\-\s": "_minus_",
    r"\s\+\s": "_plus_",
    r"\s\*\s": "_mul_",
    r"\s\/\s": "_div_"
}


def _parquet_name(stamp: datetime | None = None, *, constant: bool = False) -> Path:
    """
    Builds a parquet path inside PQ_DIR
    Args:
        stamp: Optional timestamp to use in the filename
        constant: If True, returns a constant filename (scrobble.parquet) instead of timestamped
    Returns:
        Path to the parquet file
    """
    if constant:
        return PQ_DIR / "scrobble.parquet"
    now = datetime.now(UTC)
    stamp = stamp or now
    return PQ_DIR / f"scrobbles_{stamp:%Y%m%d_%H%M%S}.parquet"


def append_or_create_parquet(df: pd.DataFrame, path: Path) -> None:
    """
    Appends data to an existing parquet file or creates a new one if it doesn't exist
    Args:
        df: DataFrame containing the data to append
        path: Path to the parquet file
    """
    import logging
    logger = logging.getLogger(__name__)
    # Column name mappings between Last.fm API and CanonFodder DB
    column_mappings = {
        # Last.fm API → CanonFodder DB
        "Artist": "artist_name",
        "Song": "track_title",
        "Album": "album_title",
        "artist mbid": "artist_mbid",
        # Ensuring we also handle the canonical names
        "artist_name": "artist_name", 
        "track_title": "track_title",
        "album_title": "album_title",
        "artist_mbid": "artist_mbid"
    }
    # Define the expected column order (matching what's used in profile.py)
    expected_column_order = ["artist_name", "album_title", "play_time", "track_title", "artist_mbid"]

    # Creating a normalized copy of the dataframe
    normalized_df = df.copy()
    # Normalizing column names in the new dataframe
    for original, normalized in column_mappings.items():
        if original in normalized_df.columns and normalized not in normalized_df.columns:
            normalized_df[normalized] = normalized_df[original]
    # Handling timestamp/play_time conversion
    if "play_time" not in normalized_df.columns and "uts" in normalized_df.columns:
        normalized_df["play_time"] = pd.to_datetime(normalized_df["uts"], unit="s", utc=True)
    if path.exists():
        # Reading existing data
        existing_df = pd.read_parquet(path)
        # Normalizing column names in the existing dataframe
        normalized_existing = existing_df.copy()
        for original, normalized in column_mappings.items():
            if original in normalized_existing.columns and normalized not in normalized_existing.columns:
                normalized_existing[normalized] = normalized_existing[original]
        # Standard deduplication columns in order of priority
        dedup_candidates = [
            # These three together form a unique key in the database
            ["artist_name", "track_title", "play_time"],
            # Fall back to just artist and track if no play_time
            ["artist_name", "track_title"],
            # Last resort - just artist name
            ["artist_name"]
        ]
        # Finding the first set of deduplication columns that exists in both dataframes
        dedup_cols = None
        for candidate_cols in dedup_candidates:
            if all(col in normalized_df.columns for col in candidate_cols) and \
               all(col in normalized_existing.columns for col in candidate_cols):
                dedup_cols = candidate_cols
                break
        if dedup_cols:
            logger.info(f"Deduplicating with columns: {dedup_cols}")
            # Combining and deduplicate using normalized columns
            combined_normalized = pd.concat([normalized_existing, normalized_df])
            # Keep the latest version of each record
            combined_normalized = combined_normalized.drop_duplicates(subset=dedup_cols, keep="last")
            # Only keep the expected columns and any additional columns that are in both dataframes
            # This prevents preserving unwanted columns like 'Album', 'Artist', 'Song', 'uts'
            expected_columns = ["artist_name", "album_title", "play_time", "track_title", "artist_mbid"]
            available_expected = [col for col in expected_columns if col in combined_normalized.columns]

            # Only include additional columns that are in both dataframes to avoid preserving unwanted columns
            additional_columns = [
                col for col in combined_normalized.columns 
                if col not in expected_columns and col not in ["Album", "Artist", "Song", "uts"]
            ]

            combined_df = combined_normalized[available_expected + additional_columns]
        else:
            logger.warning("No common deduplication columns found, concatenating without deduplication")
            combined_df = pd.concat([existing_df, normalized_df])
        # Ensure columns are in the expected order before saving
        available_expected_columns = [col for col in expected_column_order if col in combined_df.columns]
        other_columns = [col for col in combined_df.columns if col not in expected_column_order]
        final_column_order = available_expected_columns + other_columns
        combined_df = combined_df[final_column_order]
        combined_df.to_parquet(path, index=False)
        print(f"[io] parquet updated with {len(df)} rows → {path} (total: {len(combined_df)} rows)")
    else:
        # Ensure columns are in the expected order before saving
        available_expected_columns = [col for col in expected_column_order if col in normalized_df.columns]
        other_columns = [col for col in normalized_df.columns if col not in expected_column_order]
        final_column_order = available_expected_columns + other_columns
        normalized_df = normalized_df[final_column_order]

        # Just write the new data, but ensure it has normalized column names for future consistency
        normalized_df.to_parquet(path, index=False)
        print(f"[io] new parquet created with {len(df)} rows → {path}")


def dump_latest_table_to_parquet() -> None:
    """
    Materialises the newest DB scrobble table as a parquet file.
    Uses the constant filename (scrobble.parquet) instead of a timestamped one.
    """
    df_db, latest_tbl = load_scrobble_table_from_db_to_df(engine)
    if df_db is None:
        print("No scrobble table in DB – nothing to dump.")
        return
    pq_file = PQ_DIR / "scrobble.parquet"
    df_db.to_parquet(pq_file, index=False)
    print(f"Latest scrobble table persisted → {pq_file}")


def dump_parquet(df: pd.DataFrame | None = None,
                 *,
                 stamp: datetime | None = None,
                 constant: bool = True) -> Path:
    """
    Writes scrobble table df to a .parquet file
    Args:
        df: optional dataframe to dump
        stamp: overrides the timestamp in the file name (used only if constant=False)
        constant: if True, uses a constant filename rather than a timestamped one
    Returns:
        path to the written parquet file
    """
    if df is None:
        from DB.ops import load_scrobble_table_from_db_to_df
        df, _tbl = load_scrobble_table_from_db_to_df(engine)
        if df is None:
            raise RuntimeError("No scrobble table available to dump.")
    # Define the expected column order (matching what's used in profile.py)
    expected_column_order = ["artist_name", "album_title", "play_time", "track_title", "artist_mbid"]
    # Ensure columns are in the expected order before saving
    available_expected_columns = [col for col in expected_column_order if col in df.columns]
    other_columns = [col for col in df.columns if col not in expected_column_order]
    final_column_order = available_expected_columns + other_columns
    df_ordered = df[final_column_order]
    out = _parquet_name(stamp, constant=constant)
    if constant:
        append_or_create_parquet(df_ordered, out)
    else:
        # Legacy behavior - create a new timestamped file, but still ensure column order
        df_ordered.to_parquet(out, index=False)
        print(f"[io] parquet written → {out}")
    return out


def latest_parquet(*, return_df: bool = False, use_constant: bool = True):
    """
    Returns the scrobble parquet file (either constant or newest timestamped one)
    or tuple (df + path) when `return_df`
    Args:
        return_df: If True, returns a tuple (dataframe, path), otherwise just the path
        use_constant: If True, looks for scrobble.parquet, otherwise finds the newest timestamped file
    Returns:
        Path to the parquet file or tuple (DataFrame, Path) if return_df=True
    """
    if use_constant:
        constant_file = PQ_DIR / "scrobble.parquet"
        if constant_file.exists():
            if return_df:
                df = pd.read_parquet(constant_file)
                return df, constant_file
            return constant_file
    # If last.fm API fetch lacks in quality, we fall back to LB JSON fetch
    files = sorted(PQ_DIR.glob("jurda_scrobble.parquet"), reverse=True)
    if not files:
        return (None, None) if return_df else None
    newest = files[0]
    if return_df:
        df = pd.read_parquet(newest)
        return df, newest
    return newest


def load_country_timeline(path: Path) -> pd.DataFrame:
    tl = (
        pd.read_parquet(path)
        .rename(columns={"country_code": "UserCountry"})
    )
    tl["start_date"] = pd.to_datetime(tl["start_date"]).dt.normalize()
    tl["end_date"] = pd.to_datetime(tl["end_date"]).dt.normalize()
    tl.sort_values("start_date", inplace=True)
    return tl


def register_custom_palette(palette_name, palettes):
    """
    Register a custom palette in Seaborn from the palette dictionary.
    """
    import seaborn as sns
    palette = next((p for p in palettes if p["paletteName"] == palette_name), None)
    if not palette:
        raise ValueError(f"Palette {palette_name} not found in the JSON file.")
    colors = [
        f"#{color['hex']}" if not color["hex"].startswith("#") else color["hex"]
        for color in sorted(palette["colors"], key=lambda x: x["position"])
    ]
    sns.set_palette(sns.color_palette(colors))
    return colors


def sanitize(col: str, seen: Counter) -> str:
    """
    Turn 'partial_ratio - QRatio'  →  'partial_ratio_minus_QRatio'
    Guarantees each result is a valid Python identifier **and unique**.
    """
    safe = col
    for pat, tok in OP_TOKENS.items():
        safe = re.sub(pat, tok, safe)          # to encode operator
    safe = re.sub(r"\W+", "_", safe).strip("_")  # to clean leftovers
    # Guaranteeing uniqueness
    count = seen[safe]
    seen[safe] += 1
    if count:
        safe = f"{safe}_{count}"               # to append _1, _2 …
    return safe


def save_as_word_table(dataframe, file_name):
    """
    Writes `dataframe` into a Word table saved as `file_name`.
    """
    doksi = Document()
    doksi.add_heading("Categorical Feature Summary", level=1)
    table = doksi.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(dataframe.columns):
        table.cell(0, idx).text = column
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doksi.save(file_name)
