from contextlib import contextmanager
from datetime import datetime, timezone
from dateutil import parser
from dotenv import load_dotenv
load_dotenv()
from DB.models import ArtistVariantsCanonized
from docx import Document
import featuretools as ft
import fitz
from functools import partial
import getpass
from glob import glob
from hdbscan import HDBSCAN
import numpy as np
import os
import pandas as pd
from pathlib import Path
import questionary
from rapidfuzz import fuzz, process
import re
from scipy import stats
from scipy.spatial.distance import cdist
import seaborn as sns
from sklearn.feature_selection import VarianceThreshold
from sklearn.inspection import permutation_importance
from sklearn.metrics import pairwise_distances, silhouette_score
from sklearn.neighbors import KNeighborsClassifier
from sklearn.preprocessing import RobustScaler
from sqlalchemy import select
import tabulate
from typing import Any, Sequence, Tuple

HERE = Path(__file__).resolve().parent


@contextmanager
def get_db_session(*, readonly: bool = False) -> Any:
    """
    Helper that gives caller a Session and closes it afterwards.
    Imported lazily, so it never fires during package initialisation.
    """
    from DB import SessionLocal
    session = SessionLocal()
    try:
        yield session
        if not readonly:
            session.commit()
    finally:
        session.close()


def _apply_canonical(canon, variants, data, counts):
    for old in variants:
        if old != canon:
            data.loc[data["Artist"] == old, "Artist"] = canon
            counts.loc[counts["Artist"] == old, "Artist"] = canon


def _clf_scorer(x: str, y: str, clf, **kwargs) -> float:
    return clf_proba(x, y, clf)


def _remember(sig: str, name: str, sess):
    sess.execute(
        ArtistVariantsCanonized.insert().values(
            group_signature=sig,
            canonical_name=name,
            timestamp=datetime.utcnow()
        )
    )


# helper.py  ─────────────────────────────────────────────────────────────
def ask(question: str,
        default: str | None = None) -> str:
    """
    Prompt until the user types something (or presses <Enter> to accept *default*).
    Example
    -------
    answer = helper.ask("Redownload? y / n", default="n").lower()
    """
    while True:
        prompt = f"{question.strip()} "
        if default is not None:
            prompt += f"[{default}] "
        print(prompt, end='', flush=True)  # show immediately
        reply = input().strip()
        if reply:
            return reply
        if default is not None:  # empty ↦ default
            return default
        print("Please enter a value.")


def calculate_clustering_metrics(name, labels, data, cluster_centers=None, model=None):
    """
    Calculates clustering QA metrics: Noise Percentage, Silhouette Score, weighted WSS, and BIC.
    Parameters:
    - name: str, clustering experiment identifier
    - labels: array-like, cluster labels (noise points should be -1 for DBSCAN)
    - data: array-like, original dataset
    - cluster_centers: array-like, optional cluster centers (needed for WSS)
    - model: clustering model object (used for BIC if supported)
    Returns:
    - dict with clustering metrics
    """
    non_noise_indices = labels != -1
    clustered_data = data[non_noise_indices].to_numpy()
    clustered_labels = labels[non_noise_indices]
    # Calculating noise percentage
    noise_percentage = 1 - np.sum(non_noise_indices, axis=0) / len(labels)
    # Calculating silhouette score
    if len(np.unique(clustered_labels)) > 1:
        silhouette = silhouette_score(clustered_data, clustered_labels)
    else:
        silhouette = np.nan
    # Calculating WSS
    data = np.array(data)
    cluster_centers = np.array(cluster_centers)
    if cluster_centers is not None:
        wss = 0
        total_points = 0
        for i in range(len(cluster_centers)):
            # Extract points belonging to the current cluster
            cluster_points = clustered_data[clustered_labels == i]
            total_points += len(cluster_points)
            for j in range(len(cluster_points)):
                squared_distance = np.sum(
                    (cluster_points[j] - cluster_centers[i]) ** 2, axis=0
                )
                wss += squared_distance
        if total_points > 0:
            weighted_wss = wss / total_points
        else:
            weighted_wss = np.nan
    else:
        weighted_wss = np.nan
    # Calculating BIC
    if model is not None:
        n_clusters = len(np.unique(clustered_labels))
        n_features = data.shape[1]
        n_samples = len(clustered_data)
        # Log likelihood approximation for BIC
        if cluster_centers is not None:
            distances = cdist(clustered_data, cluster_centers)
            min_distances = np.min(distances, axis=1)
            log_likelihood = -0.5 * np.sum(min_distances ** 2, axis=0)
        else:
            log_likelihood = np.nan
        # BIC calculation
        n_params = n_clusters * n_features  # Approximate number of parameters
        bic = -2 * log_likelihood + n_params * np.log(n_samples)
    else:
        bic = np.nan  # BIC requires a model
    return {
        "Clustering Name": name,
        "Noise Percentage": noise_percentage,
        "Weighted WSS": weighted_wss,
        "Silhouette Score": silhouette,
        "BIC": bic,
    }


def choose_lastfm_user() -> str:
    """
    Ask once for the Last.fm user.
    ─ behaviour ─
    • If LASTFM_USER is set in .env / environment → offer it as the default
    • Empty input while no default is known → keep asking
    """
    default = os.getenv("LASTFM_USER", "").strip() or None
    while True:
        tail = f" [{default}]" if default else ""
        reply = input(f"If you are querying data for last.fm user {tail}, press enter"
                      " Otherwise, type username here: › ").strip()
        if reply:
            return reply
        if default:
            return default
        print("Please type a user name (or set LASTFM_USER in your .env).")


def clf_proba(a: str, b: str, clf) -> float:
    vec = np.fromiter(fuzzy_scores(a, b).values(), dtype=float)[None, :]
    return float(clf.predict_proba(vec)[0, 1])


def cramers_v(x, y):
    """To calculate Cramér's V for two categorical variables"""
    contingency_table = pd.crosstab(x, y)
    chi2 = stats.chi2_contingency(contingency_table)[0]
    n = contingency_table.sum().sum()
    phi2 = chi2 / n
    r, k = contingency_table.shape
    return np.sqrt(phi2 / min(r - 1, k - 1))


def fuzzy_scores(a: str, b: str) -> dict:
    return {
        "ratio": fuzz.ratio(a, b),
        "partial_ratio": fuzz.partial_ratio(a, b),
        "token_set_ratio": fuzz.token_set_ratio(a, b),
        "partial_token_ratio": fuzz.partial_token_set_ratio(a, b),
    }


def get_latest_csv(directory, prefix="jurda", extension=".csv"):
    search_pattern = os.path.join(directory, f"{prefix}_*{extension}")
    csv_files = glob(search_pattern)
    if not csv_files:
        raise FileNotFoundError(
            f"No files matching {search_pattern} found in {directory}"
        )
    latest_file = max(csv_files, key=os.path.getmtime)
    return latest_file


def load_latest_parquet_in_pq():
    """
    Scan the 'PQ' folder for files matching the pattern:
        scrobbles_YYYYMMDD_HHMMSS.parquet
    Parse out the datetime portion, identify the newest file, and load it into
    a pandas df.
    Returns (df, latest_filename) or (None, None) if none found.
    """
    pq_folder = "PQ"
    pattern = re.compile(r"^scrobbles_(\d{8}_\d{6})\.parquet$")
    files = glob(os.path.join(pq_folder, "scrobbles_*.parquet"))
    if not files:
        print("No scrobbles_*.parquet files found in PQ folder.")
        return None, None
    latest_file = None
    latest_stamp = None
    # Looping over each parquet file in PQ to find the newest
    for filepath in files:
        filename = os.path.basename(filepath)  # e.g. "scrobbles_20250416_101515.parquet"
        match = pattern.match(filename)
        if match:
            dt_str = match.group(1)  # e.g. "20250416_101515"
            try:
                dt_parsed = pd.to_datetime(dt_str, format="%Y%m%d_%H%M%S")
            except ValueError:
                continue
            if (latest_stamp is None) or (dt_parsed > latest_stamp):
                latest_stamp = dt_parsed
                latest_file = filepath
    if not latest_file:
        print("No properly named scrobbles_YYYYMMDD_HHMMSS.parquet files found.")
        return None, None
    print(f"Latest parquet determined: {os.path.basename(latest_file)}")
    # Loading the DataFrame from the newest parquet file
    df = pd.read_parquet(latest_file)
    print(f"Loaded {len(df)} rows from {latest_file}.")
    return df, latest_file


def missing_value_ratio(col):
    return (col.isnull().sum() / len(col)) * 100


def most_similar(name: str,
                 choices: Sequence[str],
                 clf,
                 threshold: float = 0.5) -> Tuple[str | None, float]:
    scorer = partial(_clf_scorer, clf=clf)
    match, score, _ = process.extractOne(
        name, choices,
        scorer=scorer,
        score_cutoff=threshold
    )
    return match, score or 0.0


def register_custom_palette(palette_name, palettes):
    """Register a custom palette in Seaborn from the palette dictionary."""
    palette = next((p for p in palettes if p["paletteName"] == palette_name), None)
    if not palette:
        raise ValueError(f"Palette {palette_name} not found in the JSON file.")
    # Extract and sanitize hex colors
    colors = [
        f"#{color['hex']}" if not color["hex"].startswith("#") else color["hex"]
        for color in sorted(palette["colors"], key=lambda x: x["position"])
    ]
    # Register the palette in Seaborn
    sns.set_palette(sns.color_palette(colors))
    print(f"Custom palette '{palette_name}' applied successfully.")
    return colors


def safe_parse(ts_raw: str | int | None):
    """
    Try to turn *ts_raw* into an aware UTC datetime.
    •  epoch seconds          → datetime(… , tzinfo=UTC)
    •  almost any date string → datetime(… , tzinfo=UTC) via dateutil
    •  unknown / in-progress  → return the original string
    """
    if ts_raw is None:
        return None
    # 1) Epoch seconds?
    try:
        return datetime.fromtimestamp(int(ts_raw), tz=timezone.utc)
    except (ValueError, TypeError, OverflowError):
        pass
    # 2) Let dateutil figure it out
    try:
        dt = parser.parse(ts_raw)
        # Making sure we always have an explicit tz
        return dt.astimezone(timezone.utc) if dt.tzinfo else dt.replace(tzinfo=timezone.utc)
    except (ValueError, TypeError):
        return ts_raw  # last-resort


def save_as_word_table(dataframe, file_name):
    doksi = Document()
    doksi.add_heading("Categorical Feature Summary", level=1)
    table = doksi.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(dataframe.columns):
        table.cell(0, idx).text = column
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doksi.save(file_name)


# Creating shortened names for artists
short_artists = {
    "": "",
    "": "",
    "": ""
}


def unify_artist_names_cli(
        data,
        fltrd_artcount,
        similar_artist_groups,
):
    """
    identical doc-string – but we talk SQLAlchemy now, not sqlite3.
    """
    from DB import get_engine as _get_engine, get_session as _get_session
    with _get_session().begin() as sess:
        for group in similar_artist_groups:
            group = list(group)
            if len(group) <= 1:
                continue
            sig = "|".join(sorted(group))
            existing = sess.scalar(
                select(ArtistVariantsCanonized.canonical_name)
                .where(ArtistVariantsCanonized.group_signature == sig)
            )
            # -----------------------------------------------------------------
            if existing is not None:  # handled earlier
                if existing == "__SKIP__":
                    print(f"\nUser previously SKIPPED group {group}.")
                    continue
                print(f"\nUser previously unified {group} to "
                      f"'{existing}'. Applying again.")
                _apply_canonical(existing, group, data, fltrd_artcount)
                continue
            # ----------------------------- ask the user ----------------------
            print("\n---")
            print(f"These artist names appear to be duplicates:\n{group}")
            choice = questionary.select(
                "Which name would you like to keep for all occurrences?",
                choices=group + ["Custom", "Skip"]
            ).ask()
            if not choice or choice == "Skip":
                _remember(sig, "__SKIP__", sess)
                continue
            canonical = (questionary.text("Enter a custom canonical name:").ask()
                         if choice == "Custom" else choice)
            if not canonical:
                _remember(sig, "__SKIP__", sess)
                continue
            _apply_canonical(canonical, group, data, fltrd_artcount)
            _remember(sig, canonical, sess)
    # ---------- re-aggregate counts ------------------------------------------
    out = data["Artist"].value_counts().reset_index(names=["Artist", "Count"])
    return data, out


def variance_testing(dframe, varthresh):
    selector = VarianceThreshold(threshold=varthresh)
    _ = selector.fit_transform(dframe)
    variances = selector.variances_
    variance_df = pd.DataFrame({"features": dframe.columns, "variances": variances})
    variance_df = variance_df.sort_values(by="variances", ascending=False)
    selected_features = dframe.columns[selector.get_support(indices=True)]
    return variance_df, selected_features


def verify_volcano_name(csv_path: str | Path) -> None:
    """
    Show whether the artist column contains either spelling of
    'Volcano(,) I'm Still Excited!!'  (comma‑preserved check).
    """
    df = pd.read_csv(csv_path, dtype=str, keep_default_na=False)
    # First column is artist
    artist_col = df.columns[0]
    with_comma = df[artist_col].eq("Volcano, I'm Still Excited!!").sum()
    no_comma = df[artist_col].eq("Volcano I'm Still Excited!!").sum()
    print(f"\n── Let's check if CSV export preserves commas. ──")
    print(f"\n── We are querying the band name Volcano, I'm Still Excited!!. ──")
    print(f"Rows with *comma*  : {with_comma}")
    print(f"Rows without comma : {no_comma}")
    print("─────────────────────\n")


def winsorization_outliers(df):
    out = []
    for n in df:
        q1 = np.percentile(df, 1)
        q3 = np.percentile(df, 99)
        if n > q3 or n < q1:
            out.append(n)
    print("Outliers:", out)
    return out


def yes_no(question: str, *, default: str = "n") -> bool:
    return ask(question, default).lower().startswith("y")
